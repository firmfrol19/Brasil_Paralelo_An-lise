import docx
import pandas as pd
import matplotlib.pyplot as plt
#matplotlib inline
import seaborn as sns

doc = docx.Document('tabela2.docx')

try:
  document = docx.Document('tabela2.docx')
except:
  document = docx.Document()
  document.save('tabela2.docx')
  print("Previous file was corrupted or didn't exist - new file was created.")

count = 0 
for table in doc.tables:
  print(table)
  for row in table.rows:
    count += 1
    for cell in row.cells:
      for para in cell.paragraphs:
        print(para.text)
        print(count)

documento = docx.Document('tabela2.docx')

table = documento.tables[0]
print(table)

data = []
keys = None
for i, row in enumerate(table.rows):
    text = (cell.text for cell in row.cells)
keys = None
for i, row in enumerate(table.rows):
    text = (cell.text for cell in row.cells)
    print(i, row)

    if i == 0:
      keys = tuple(text)
      continue
    row_data = dict(zip(keys, text))
    if row_data not in data:
      data.append(row_data)
    else:
      break

tbl_01 = pd.DataFrame(data=data)
tbl_01
print(tbl_01)

Categoria = []
Quantidade = []
for dicionário in data:
  for k, v in dicionário.items():
    if len(k) == 9:
      if k not in Categoria:
        Categoria.append(v)
    else:
      if k not in Quantidade:
        Quantidade.append(v)
print(Categoria, Quantidade)
lista_categoria = Categoria
print(lista_categoria)

lista_quantidade = []
for n in Quantidade:
    na = int(n)
    lista_quantidade.append(na)

#print("ESSE É O CUMPRIMENTO DAS LISTAS:", len(lista_categoria), len(lista_quantidade))
#print("ESSA É A LISTA QUANTIDADE:", len(Quantidade))
#print("Essa é a lista_quantidade:", len(lista_quantidade))


indice = []
for i in range(len(lista_categoria)):
    indice.append(i+1)

# Gráfico 2 
dimensao = (12, 7)
fig, ax = plt.subplots(figsize=dimensao)
g = sns.barplot(x= indice, y= lista_quantidade,  orient = "v",  palette = "inferno_r", ci = None)
# Consultar palette de cores no site: https://medium.com/@morganjonesartist/color-guide-to-seaborn-palettes-da849406d44f (acesso no dia 11/11/2020)
g.set(xlabel = 'Categoria', ylabel= 'Frequência')
plt.title("""Frequência de Categorias - Documentário: 'Brasil: Última Cruzada' (2017-2018)""", fontsize=16)
fig.tight_layout()
#plt.savefig("grafico2_1.png")
plt.show()



